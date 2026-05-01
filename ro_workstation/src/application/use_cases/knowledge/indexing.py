from __future__ import annotations

import hashlib
from datetime import UTC, datetime
from pathlib import Path

from docx import Document

from src.core.paths import project_path
from src.domain.schemas.knowledge import IndexedDocument
from src.infrastructure.llm.embeddings import get_embedder
from src.infrastructure.persistence.json_repo import JsonRepository


SUPPORTED_EXTENSIONS = {".txt", ".md", ".docx"}


class KnowledgeIndexingService:
    def __init__(self) -> None:
        self.knowledge_dir = project_path("data", "knowledge_docs")
        self.knowledge_dir.mkdir(parents=True, exist_ok=True)
        self.registry = JsonRepository(project_path("data", "knowledge_index.json"), [])
        self.client = self._create_client()

    def list_documents(self) -> list[IndexedDocument]:
        return [IndexedDocument.model_validate(item) for item in self.registry.read()]

    def store_uploaded_documents(self, uploaded_files) -> list[Path]:
        saved_paths = []
        for uploaded_file in uploaded_files:
            target = self.knowledge_dir / uploaded_file.name
            target.write_bytes(uploaded_file.getbuffer())
            saved_paths.append(target)
        return saved_paths

    @staticmethod
    def _create_client():
        try:
            import chromadb

            return chromadb.PersistentClient(path=str(project_path("data", "chroma_db")))
        except Exception:
            class _FallbackCollection:
                def upsert(self, **kwargs):
                    return None

            class _FallbackClient:
                def get_or_create_collection(self, name):
                    return _FallbackCollection()

            return _FallbackClient()

    def index_uploaded_documents(self, uploaded_files, department: str, uploaded_by: str) -> list[dict]:
        saved_paths = self.store_uploaded_documents(uploaded_files)
        return [self.index_saved_document(path, department, uploaded_by).model_dump() for path in saved_paths]

    def index_saved_document(self, path: Path, department: str, uploaded_by: str) -> IndexedDocument:
        text = self._extract_text(path)
        chunks = self._chunk_text(text)
        if not chunks:
            raise ValueError("The uploaded document does not contain readable text.")

        embedder = get_embedder()
        embeddings = embedder.encode(chunks)
        if hasattr(embeddings, "tolist"):
            embeddings = embeddings.tolist()

        collection = self.client.get_or_create_collection(name="bank_knowledge")
        source_name = path.name
        ids = [hashlib.md5(f"{source_name}:{idx}:{chunk}".encode("utf-8")).hexdigest() for idx, chunk in enumerate(chunks)]
        metadata = [
            {
                "title": path.stem,
                "department": department or "ALL",
                "uploaded_by": uploaded_by or "unknown",
                "source_file": source_name,
                "chunk_index": idx,
            }
            for idx in range(len(chunks))
        ]
        collection.upsert(ids=ids, documents=chunks, embeddings=embeddings, metadatas=metadata)

        indexed = IndexedDocument(
            file_name=path.name,
            department=department or "ALL",
            uploaded_by=uploaded_by or "unknown",
            chunks=len(chunks),
            indexed_at=datetime.now(UTC),
        )
        records = [item.model_dump(mode="json") for item in self.list_documents() if item.file_name != path.name]
        records.append(indexed.model_dump(mode="json"))
        self.registry.write(records)
        return indexed

    @staticmethod
    def _chunk_text(text: str, chunk_size: int = 1200) -> list[str]:
        cleaned = " ".join(text.split())
        return [cleaned[index:index + chunk_size] for index in range(0, len(cleaned), chunk_size)] if cleaned else []

    @staticmethod
    def _extract_text(path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8", errors="ignore")
        if suffix == ".docx":
            doc = Document(str(path))
            return "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        raise ValueError(f"Unsupported file type: {suffix}. Supported types are {', '.join(sorted(SUPPORTED_EXTENSIONS))}.")
